import fitz

from os import listdir, remove
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont


class Converter:
    def __init__(self, leave_number_in_paragraph_title: bool = True):

        self.leave_number_in_paragraph_title = leave_number_in_paragraph_title

        self.document = None
        self.dpi = 200  # dots per inch
        self.font = ImageFont.truetype('arial.ttf', 40)

    def __prepare_to_convert(self):
        self.__prepare_document()
        self.__prepare_variables()

    def __prepare_document(self):
        self.document = Document()

        section = self.document.sections[0]
        margin = Mm(0)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    def __prepare_variables(self):
        self.queue_images = []
        self.counter = 0

    def __draw_extra_information(self, image_filename: str, paragraph_title=None):
        black = (0, 0, 0, 255)

        image = Image.open(image_filename)
        draw = ImageDraw.Draw(image)

        page = str(self.counter).zfill(3)

        desired_size = 1654, 2339
        text_y = image.height * 2200 / desired_size[1]

        is_even = self.counter % 2
        if is_even:
            page_text_pos = [image.width * 1500 / desired_size[0], text_y]
        else:
            page_text_pos = [image.width * 130 / desired_size[0], text_y]
        draw.text(page_text_pos, page, font=self.font, fill=black)

        if paragraph_title:
            filename_text_pos = [270, text_y]
            draw.text(filename_text_pos, paragraph_title, font=self.font, fill=black)

        image.save(image_filename)

    def __make_sheet(self):
        """
            Сначала добавляет все элементы, которые должны находится на чётных страницах,
        а потом остальные. То есть, сначала 1, 3, 5, 7 страницу на лицевую сторону, а потом
        2, 4, 6, 8 на обратную, чтоб корректно выполнялась двухсторонняя печать
        """

        for is_back_side in [False, True]:
            table = self.__add_table()

            start_from = 1 if is_back_side else 0
            for image_n in range(start_from, len(self.queue_images), 2):
                cell_n = (image_n - int(is_back_side)) // 2

                cells_in_row = table.rows[cell_n // 2].cells
                column_n = cell_n % 2

                if is_back_side:
                    """
                    На обратной стороне файлы начинают добавляться с обратной стороны.
                    
                    Например:
                        - 1 страница будет слева сверху на листе
                        - 3 уже должна быть справа сверху (на обратной стороне)
                        
                    Чтобы опять же была возможна двухстороняя печать.
                    """
                    column_n = -column_n - 1

                table_cell = cells_in_row[column_n].paragraphs[0]
                if not is_back_side:
                    table_cell.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                run = table_cell.add_run()
                image = self.queue_images[image_n]
                image_width = Mm(95)
                r = run.add_picture(image, width=image_width)
                remove(image)

            self.document.add_page_break()

        self.queue_images.clear()

    def __add_table(self):
        table = self.document.add_table(rows=2, cols=2)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        return table

    def __get_service_symbols(self):
        if self.leave_number_in_paragraph_title:
            return '.pdf'
        else:
            return '0123456789.pdf'

    def execute(self, out_filename: str, folder_with_pdfs_path: str = None, pdf_path: str = None):
        if bool(folder_with_pdfs_path) == bool(pdf_path):
            return

        pdf_files = []
        if folder_with_pdfs_path:
            folder_with_pdfs_path = folder_with_pdfs_path.strip('/\\')
            pdf_files = listdir(folder_with_pdfs_path)
        elif pdf_path:
            pdf_files = [pdf_path]

        self.__prepare_to_convert()

        for pdf_file_index, pdf_filename in enumerate(pdf_files):
            if folder_with_pdfs_path:
                file_to_convert = f'{folder_with_pdfs_path}/{pdf_filename}'
            elif pdf_path:
                file_to_convert = pdf_filename

            pdf_file = fitz.open(file_to_convert)
            images = []
            for num, page in enumerate(pdf_file.pages()):
                pixmap = page.get_pixmap(dpi=self.dpi)
                images.append(pixmap)
            pdf_file.close()

            if len(pdf_files) > 1:
                paragraph_title_max_length = 40
                service_symbols = self.__get_service_symbols()
                paragraph_title = pdf_filename.strip(service_symbols)
                paragraph_title = paragraph_title[:paragraph_title_max_length]
            else:
                paragraph_title = None

            for image_index, image in enumerate(images):
                image_filename = str(self.counter) + '.png'
                image.save(image_filename, 'PNG')
                self.counter += 1

                self.__draw_extra_information(image_filename, paragraph_title=paragraph_title)
                self.queue_images.append(image_filename)

                is_images_queue_full = len(self.queue_images) == 8
                is_last_pdf = pdf_file_index == len(pdf_files) - 1
                is_last_page_in_pdf = image_index == len(images) - 1

                if is_images_queue_full or (is_last_pdf and is_last_page_in_pdf):
                    self.__make_sheet()

        self.document.save(out_filename)
